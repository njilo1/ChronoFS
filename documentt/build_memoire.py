# -*- coding: utf-8 -*-
"""Assemble le memoire ChronoFS (.docx) a partir de memoire_contenu.md,
en respectant la mise en forme de la charte des memoires FS-UEB.

Usage: python build_memoire.py
Sortie: ChronoFS_Memoire.docx (dans ce meme dossier)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
ASSETS = HERE / "memoire_assets"
LOGO_UEB = HERE.parent / "docs" / "assets" / "logo_ueb.png"
LOGO_FS = HERE.parent / "docs" / "assets" / "logo_fs.png"
CONTENU_MD = HERE / "memoire_contenu.md"
OUT_PATH = HERE / "ChronoFS_Memoire.docx"

TITRE_MEMOIRE = ("CONCEPTION ET RÉALISATION D'UNE APPLICATION WEB DE GÉNÉRATION "
                  "AUTOMATIQUE DES EMPLOIS DU TEMPS : CAS DE LA FACULTÉ DES SCIENCES "
                  "DE L'UNIVERSITÉ D'ÉBOLOWA")
AUTEUR = "TCHAMBA NJILO FERDINAND"
MATRICULE = "23I0076FS"
ENCADREUR_1 = "Dr KENGNI OLGA"
GRADE_ENCADREUR_1 = "Chargée de Cours"
ENCADREUR_2 = "Dr NYABEYE DORIS"
GRADE_ENCADREUR_2 = "Assistante"
ANNEE_ACAD = "2025 - 2026"

ROMAN_CHAP = {1: "I", 2: "II", 3: "III"}

# ---------------------------------------------------------------------------
# Helpers bas niveau (OOXML)
# ---------------------------------------------------------------------------

def set_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def justify(paragraph, spacing15=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    if spacing15:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)


def add_field(paragraph, instr, result_text="1"):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = result_text
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld1); r.append(instrText); r.append(fld2); r.append(t); r.append(fld3)
    set_font(run, size=12)
    return run


def set_page_number_format(section, fmt, start):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))


def set_margins(section, cm=2.5):
    section.top_margin = Cm(cm)
    section.bottom_margin = Cm(cm)
    section.left_margin = Cm(cm)
    section.right_margin = Cm(cm)


def new_section(doc, unlink_footer=True):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(sec)
    if unlink_footer:
        sec.footer.is_linked_to_previous = False
        for p in list(sec.footer.paragraphs):
            p.clear()
    return sec


def set_footer_page_number(section, align=WD_ALIGN_PARAGRAPH.RIGHT):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = align
    add_field(p, "PAGE")


def set_title_page(section, enable=True):
    """Active w:titlePg : la 1ere page de la section a son propre pied de
    page (utilise pour masquer le numero sur la page de titre de chapitre,
    conformement a la charte, sans casser la continuite de la numerotation)."""
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if enable and titlePg is None:
        sectPr.append(OxmlElement('w:titlePg'))
    elif not enable and titlePg is not None:
        sectPr.remove(titlePg)


def start_chapter_section(doc):
    """Nouvelle section pour un chapitre : la 1ere page (titre du chapitre)
    n'affiche pas de numero de page, mais la numerotation continue reste
    inchangee pour les pages suivantes (charte : page de garde de chapitre
    non numerotee)."""
    sec = new_section(doc, unlink_footer=True)
    set_title_page(sec, True)
    set_footer_page_number(sec)
    fp_footer = sec.first_page_footer
    fp_footer.is_linked_to_previous = False
    for p in list(fp_footer.paragraphs):
        p.clear()
    return sec


def enable_update_fields_on_open(doc):
    element = doc.settings.element
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    element.append(upd)


# ---------------------------------------------------------------------------
# Compteurs de figures / tableaux (numerotation manuelle de secours,
# les champs SEQ se remettront a jour dans Word via F9 / mise a jour auto)
# ---------------------------------------------------------------------------
fig_counter = 0
tab_counter = 0


def _style_caption_run(run):
    set_font(run, size=11, bold=True, italic=True)
    run.font.underline = True


def add_figure(doc, filename, caption, width_cm=14):
    global fig_counter
    fig_counter += 1
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r0 = cap.add_run("Figure "); _style_caption_run(r0)
    rseq = add_field(cap, 'SEQ Figure \\* ARABIC', str(fig_counter)); _style_caption_run(rseq)
    r1 = cap.add_run(" : "); _style_caption_run(r1)
    r2 = cap.add_run(caption); _style_caption_run(r2)
    return fig_counter


def add_table_caption(doc, title):
    global tab_counter
    tab_counter += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    r0 = cap.add_run("Tableau "); _style_caption_run(r0)
    rseq = add_field(cap, 'SEQ Tableau \\* ARABIC', str(tab_counter)); _style_caption_run(rseq)
    r1 = cap.add_run(" : "); _style_caption_run(r1)
    r2 = cap.add_run(title); _style_caption_run(r2)
    return tab_counter


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    return table


# ---------------------------------------------------------------------------
# Rendu de paragraphe avec **gras** inline
# ---------------------------------------------------------------------------

def add_rich_paragraph(doc, text, size=12, align_justify=True, italic=False):
    p = doc.add_paragraph()
    if align_justify:
        justify(p)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, size=size, bold=True, italic=italic)
        else:
            r = p.add_run(part)
            set_font(r, size=size, bold=False, italic=italic)
    return p


def add_heading(doc, text, level=1):
    """Utilise les styles Word natifs Heading 1/2/3 (configures en TNR par
    configure_heading_styles) afin que la table des matieres (champ TOC)
    detecte automatiquement ces titres."""
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p


def add_chapter_divider(doc, roman, title):
    """Page de garde de chapitre (non numerotee dans le flux visuel).
    Un seul paragraphe Heading 1 (pour la table des matieres), mis en
    valeur visuellement en grand et centre sur sa propre page."""
    for _ in range(6):
        doc.add_paragraph()
    p1 = add_heading(doc, f"CHAPITRE {roman} : {title.upper()}", level=1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p1.runs:
        set_font(r, size=18, bold=True)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Parsing du contenu markdown maison
# ---------------------------------------------------------------------------

def parse_blocks(md_text):
    lines = md_text.splitlines()
    blocks = []  # list of (title, body_lines)
    current_title = None
    current_body = []
    for line in lines:
        if line.startswith("## "):
            if current_title is not None:
                blocks.append((current_title, current_body))
            current_title = line[3:].strip()
            current_body = []
        elif line.startswith("# ") or line.strip() == "---" or line.startswith(">"):
            continue
        else:
            current_body.append(line)
    if current_title is not None:
        blocks.append((current_title, current_body))
    return blocks


def add_compact_body(doc, body_lines, size=11):
    """Paragraphes en interligne simple et espacement reduit, pour le bloc
    Resume/Abstract que la charte limite a une page."""
    for line in body_lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); set_font(r, size=size, bold=True)
            else:
                r = p.add_run(part); set_font(r, size=size, bold=False)


def render_body(doc, body_lines, chapter_num=None):
    i = 0
    n = len(body_lines)
    while i < n:
        line = body_lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # Heading level 2 (### N.M. Title)
        m = re.match(r"^### (\d+)\.(\d+)\.\s+(.+)$", line)
        if m:
            c, s, title = m.groups()
            label = f"{ROMAN_CHAP.get(chapter_num, c)}.{s}. {title}" if chapter_num else f"{c}.{s}. {title}"
            add_heading(doc, label, level=2)
            i += 1
            continue
        # Heading level 3 (#### N.M.K. Title)
        m = re.match(r"^#### (\d+)\.(\d+)\.(\d+)\.\s+(.+)$", line)
        if m:
            c, s, sub, title = m.groups()
            label = f"{ROMAN_CHAP.get(chapter_num, c)}.{s}.{sub}. {title}" if chapter_num else f"{c}.{s}.{sub}. {title}"
            add_heading(doc, label, level=3)
            i += 1
            continue
        # Figure marker
        m = re.match(r"^!\[FIGURE:([^|]+)\|(.+)\]$", line)
        if m:
            fname, caption = m.groups()
            add_figure(doc, fname.strip(), caption.strip())
            i += 1
            continue
        # Table caption marker
        m = re.match(r"^!\[TABLECAPTION:(.+)\]$", line)
        if m:
            title = m.group(1).strip()
            add_table_caption(doc, title)
            i += 1
            continue
        # Markdown table
        if line.startswith("|"):
            table_lines = []
            while i < n and body_lines[i].strip().startswith("|"):
                table_lines.append(body_lines[i].strip())
                i += 1
            header = [c.strip() for c in table_lines[0].strip("|").split("|")]
            rows = []
            for tl in table_lines[2:]:
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            add_table(doc, header, rows)
            continue
        # Bullet list
        if line.startswith("- "):
            bullet_items = []
            while i < n and body_lines[i].startswith("- "):
                bullet_items.append(body_lines[i][2:].strip())
                i += 1
            for item in bullet_items:
                p = doc.add_paragraph(style='List Bullet')
                justify(p)
                parts = re.split(r"(\*\*.*?\*\*)", item)
                for part in parts:
                    if not part:
                        continue
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2]); set_font(r, size=12, bold=True)
                    else:
                        r = p.add_run(part); set_font(r, size=12)
            continue
        # Paragraph (accumulate wrapped single logical line; our source uses one line per paragraph)
        add_rich_paragraph(doc, line.strip())
        i += 1


# ---------------------------------------------------------------------------
# Pages liminaires specifiques (non markdown)
# ---------------------------------------------------------------------------

def _cell_center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cell


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _tight(paragraph, space_before=0, space_after=6):
    """Espacement maitrise pour la page de garde (sinon l'interligne 1.5
    global du style Normal fait deborder la couverture sur 2 pages)."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return paragraph


def _cp(doc, text="", size=11, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p, space_before, space_after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def _keep_table_together(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)


def add_cover_page(doc):
    # --- En-tete bilingue FR / logos / EN, sur 3 colonnes sans bordures ---
    header = doc.add_table(rows=1, cols=3)
    _no_borders(header)
    _keep_table_together(header)
    header.autofit = True
    fr_cell, logo_cell, en_cell = header.rows[0].cells

    fr_lines = ["RÉPUBLIQUE DU CAMEROUN", "Paix - Travail - Patrie",
                "MINISTÈRE DE L'ENSEIGNEMENT SUPÉRIEUR",
                "UNIVERSITÉ D'ÉBOLOWA", "FACULTÉ DES SCIENCES"]
    en_lines = ["REPUBLIC OF CAMEROON", "Peace - Work - Fatherland",
                "MINISTRY OF HIGHER EDUCATION",
                "THE UNIVERSITY OF EBOLOWA", "FACULTY OF SCIENCE"]

    fr_cell.text = ""
    for j, line in enumerate(fr_lines):
        p = fr_cell.paragraphs[0] if j == 0 else fr_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight(p, 0, 3)
        r = p.add_run(line); set_font(r, size=9, bold=False)

    en_cell.text = ""
    for j, line in enumerate(en_lines):
        p = en_cell.paragraphs[0] if j == 0 else en_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight(p, 0, 3)
        r = p.add_run(line); set_font(r, size=9, bold=False)

    logo_cell.text = ""
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(lp, 0, 2)
    if LOGO_UEB.exists():
        r = lp.add_run(); r.add_picture(str(LOGO_UEB), width=Cm(1.9))
    lp2 = logo_cell.add_paragraph(); lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(lp2, 0, 0)
    if LOGO_FS.exists():
        r2 = lp2.add_run(); r2.add_picture(str(LOGO_FS), width=Cm(1.9))

    _cp(doc, space_after=10)
    _cp(doc, "DÉPARTEMENT DES TECHNOLOGIES DE L'INFORMATION ET DE LA COMMUNICATION", size=11, bold=True, space_after=2)
    _cp(doc, "DEPARTMENT OF INFORMATION AND COMMUNICATION TECHNOLOGIES", size=10, bold=True, italic=True, space_after=24)

    _cp(doc, "MÉMOIRE", size=16, bold=True, space_after=6)
    _cp(doc, "présenté en vue de l'obtention de la Licence en Architecture des Logiciels", size=11, italic=True, space_after=24)

    # --- Titre du memoire, sobre : cadre noir simple (pas de couleur) ---
    box = doc.add_table(rows=1, cols=1)
    _keep_table_together(box)
    box.rows[0].cells[0].text = ""
    p6 = box.rows[0].cells[0].paragraphs[0]
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p6, 6, 6)
    r6 = p6.add_run(TITRE_MEMOIRE); set_font(r6, size=14, bold=True)
    cell_tcPr = box.rows[0].cells[0]._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '8'); el.set(qn('w:color'), '000000')
        borders.append(el)
    cell_tcPr.append(borders)

    _cp(doc, space_after=24)

    _cp(doc, "Par", size=11, space_after=4)
    _cp(doc, AUTEUR, size=13, bold=True, space_after=4)
    p9 = _cp(doc, space_after=20)
    r9a = p9.add_run("Matricule : "); set_font(r9a, size=10, italic=True)
    r9b = p9.add_run(MATRICULE); set_font(r9b, size=10, bold=True)

    _cp(doc, "Sous la direction de", size=11, space_after=8)

    # --- Encadreurs en deux colonnes, comme le modele ---
    enc = doc.add_table(rows=2, cols=2)
    _no_borders(enc)
    _keep_table_together(enc)
    enc.rows[0].cells[0].text = ""; enc.rows[0].cells[1].text = ""
    for cell, name in ((enc.rows[0].cells[0], ENCADREUR_1), (enc.rows[0].cells[1], ENCADREUR_2)):
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _tight(p, 0, 2)
        r = p.add_run(name); set_font(r, size=11, bold=True)
    enc.rows[1].cells[0].text = ""; enc.rows[1].cells[1].text = ""
    for cell, grade in ((enc.rows[1].cells[0], GRADE_ENCADREUR_1), (enc.rows[1].cells[1], GRADE_ENCADREUR_2)):
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _tight(p, 0, 0)
        r = p.add_run(grade); set_font(r, size=10, italic=True)

    _cp(doc, space_after=24)
    _cp(doc, f"Année académique : {ANNEE_ACAD}", size=11, bold=True)


def add_toc_page(doc):
    add_heading(doc, "TABLE DES MATIÈRES", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u',
              result_text="Sommaire à générer : clic droit → Mettre à jour les champs (ou F9).")


def add_list_figures_page(doc):
    add_heading(doc, "LISTE DES FIGURES", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\c "Figure"',
              result_text="Liste à générer : clic droit → Mettre à jour les champs (ou F9).")


def add_list_tableaux_page(doc):
    add_heading(doc, "LISTE DES TABLEAUX", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\c "Tableau"',
              result_text="Liste à générer : clic droit → Mettre à jour les champs (ou F9).")


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------

def set_base_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_heading_styles(doc):
    """Style les Heading 1/2/3 natifs de Word en Times New Roman noir (pas
    le bleu par defaut), afin qu'ils servent a la fois de rendu visuel et
    de source pour le champ TOC (table des matieres automatique)."""
    specs = {'Heading 1': 16, 'Heading 2': 14, 'Heading 3': 13}
    for name, size in specs.items():
        style = doc.styles[name]
        style.font.name = 'Times New Roman'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        pf = style.paragraph_format
        pf.space_before = Pt(14)
        pf.space_after = Pt(8)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.keep_with_next = True
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def main():
    doc = Document()
    set_base_styles(doc)
    configure_heading_styles(doc)

    # --- Section 0 : couverture(s), sans numero de page ---
    sec0 = doc.sections[0]
    set_margins(sec0)
    sec0.footer.is_linked_to_previous = False
    for p in list(sec0.footer.paragraphs):
        p.clear()

    add_cover_page(doc)
    doc.add_page_break()
    add_cover_page(doc)  # page de couverture reprise en interne (papier normal)

    # --- Section 1 : pages liminaires, pagination romaine minuscule ---
    sec1 = new_section(doc, unlink_footer=True)
    set_page_number_format(sec1, 'lowerRoman', 1)
    set_footer_page_number(sec1)

    blocks = parse_blocks(CONTENU_MD.read_text(encoding='utf-8'))
    block_map = {title: body for title, body in blocks}

    add_heading(doc, "REMERCIEMENTS", level=1)
    render_body(doc, block_map["REMERCIEMENTS"])
    doc.add_page_break()

    add_toc_page(doc)
    doc.add_page_break()
    add_list_figures_page(doc)
    doc.add_page_break()
    add_list_tableaux_page(doc)
    doc.add_page_break()

    add_heading(doc, "LISTE DES ABRÉVIATIONS, SIGLES ET ACRONYMES", level=1)
    render_body(doc, block_map["LISTE DES ABRÉVIATIONS, SIGLES ET ACRONYMES"])
    doc.add_page_break()

    add_heading(doc, "LISTE DES DÉFINITIONS", level=1)
    render_body(doc, block_map["LISTE DES DÉFINITIONS"])
    doc.add_page_break()

    # Resume + Abstract : la charte impose "01 page maximum et en un bloc"
    # -> interligne resserre (simple) specifiquement pour ce bloc.
    h_res = add_heading(doc, "RÉSUMÉ", level=1)
    h_res.paragraph_format.space_before = Pt(0)
    add_compact_body(doc, block_map["RÉSUMÉ"])
    h_abs = add_heading(doc, "ABSTRACT", level=1)
    h_abs.paragraph_format.space_before = Pt(10)
    add_compact_body(doc, block_map["ABSTRACT"])

    # --- Section 2 : corps du document, pagination arabe a partir de 1 ---
    sec2 = new_section(doc, unlink_footer=True)
    set_page_number_format(sec2, 'decimal', 1)
    set_footer_page_number(sec2)

    add_heading(doc, "INTRODUCTION GÉNÉRALE", level=1)
    render_body(doc, block_map["INTRODUCTION GÉNÉRALE"])

    for title, body in blocks:
        m = re.match(r"^CHAPITRE (\d+) — (.+)$", title)
        if not m:
            continue
        num = int(m.group(1))
        chap_title = m.group(2)
        roman = ROMAN_CHAP.get(num, str(num))
        start_chapter_section(doc)
        add_chapter_divider(doc, roman, chap_title)
        render_body(doc, body, chapter_num=num)

    doc.add_page_break()
    add_heading(doc, "CONCLUSION GÉNÉRALE, PERSPECTIVES ET RECOMMANDATIONS", level=1)
    render_body(doc, block_map["CONCLUSION GÉNÉRALE, PERSPECTIVES ET RECOMMANDATIONS"])
    doc.add_page_break()

    add_heading(doc, "RÉFÉRENCES BIBLIOGRAPHIQUES", level=1)
    for line in block_map["RÉFÉRENCES BIBLIOGRAPHIQUES"]:
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                set_font(r, size=12)

    enable_update_fields_on_open(doc)
    doc.save(str(OUT_PATH))
    print(f"OK -> {OUT_PATH}")


if __name__ == "__main__":
    main()

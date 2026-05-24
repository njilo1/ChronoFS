"""
Génération du planning au format Word (.docx) via python-docx.

Version simplifiée du PDF : en-tête bilingue allégé, tableau jours ×
créneaux par salle, NB et signature à la fin. Pas d'image cachet
(impossible de positionner aussi finement qu'en CSS sans templates Word
existants — on garde "Le Doyen" en texte).
"""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from django.conf import settings
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from core.constants import HORAIRES_CRENEAUX, Jour
from core.models import Salle, Seance, Semaine


_PAUSES_HORAIRES = [
    ('10:00', '10:15'),
    ('12:45', '13:00'),
    ('15:30', '15:45'),
]


def _set_cell_bg(cell, hex_color: str) -> None:
    """Applique un fond de couleur à une cellule de tableau Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


_ORDRE_CAMPUS = {
    'Campus Principal FS': (0, 0),
    'Lycée Classique':     (0, 1),
    'Face CRA':            (0, 2),
    'Campus Monatélé':     (1, 0),
}

def _ordre_salle(salle: Salle) -> tuple:
    return _ORDRE_CAMPUS.get(salle.campus.nom, (9, 9)) + (salle.nom,)


def _hum_horaire(deb: str, fin: str) -> str:
    def hum(h):
        h, m = h.split(':')
        return f'{int(h)}h{m}'
    return f'{hum(deb)} - {hum(fin)}'


# ── Point d'entrée public ────────────────────────────────────────────────────
def generate_planning_docx(semaine: Semaine) -> tuple[bytes, str]:
    """Renvoie (contenu_docx_bytes, nom_fichier_suggere)."""

    doc = Document()

    # Paysage A4 + marges 1.5 cm
    for section in doc.sections:
        section.orientation       = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin   = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # ── En-tête bilingue (3 colonnes via un tableau) ────────────────────────
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = True
    cell_fr, cell_ctr, cell_en = header_table.rows[0].cells

    _ecrire_bloc_fr(cell_fr, semaine)
    _ecrire_bloc_logos(cell_ctr)
    _ecrire_bloc_en(cell_en)

    doc.add_paragraph()  # espace
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PLANNING DES COURS')
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'ANNÉE ACADÉMIQUE {semaine.annee_academique.libelle}')
    run.bold = True
    run.font.size = Pt(11)

    # ── Une section par salle ───────────────────────────────────────────────
    seances = list(
        Seance.objects.filter(semaine=semaine)
        .select_related('filiere', 'ue', 'enseignant', 'salle__campus')
    )
    salles_utilisees = sorted({s.salle for s in seances}, key=_ordre_salle)

    if not salles_utilisees:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Aucune séance planifiée pour cette semaine.')
        run.italic = True
    else:
        for i, salle in enumerate(salles_utilisees):
            if i > 0:
                doc.add_page_break()

            # Bandeau métadonnées
            meta = doc.add_paragraph()
            meta.add_run(f'Semestre : {semaine.semestre}   ').bold = True
            meta.add_run(f'Campus : {salle.campus.nom}   ').bold = True
            meta.add_run(f'Salle : {salle.nom}   ').bold = True
            meta.add_run(
                f'Semaine : Du {semaine.date_debut:%d/%m/%Y} '
                f'au {semaine.date_fin:%d/%m/%Y}'
            ).bold = True

            # Tableau jours × créneaux
            seances_salle = [s for s in seances if s.salle_id == salle.id]
            _construire_tableau(doc, seances_salle)

    # ── Bloc NB + signature ─────────────────────────────────────────────────
    doc.add_paragraph()
    nb = doc.add_paragraph()
    nb_run = nb.add_run('NB :')
    nb_run.bold = True; nb_run.italic = True
    nb.add_run(' ').italic = True
    nb.add_run('1.').bold = True
    nb.add_run(
        " Ce calendrier est dynamique et susceptible d'être modifié. Les "
        "étudiants sont invités à bien vouloir régulièrement (quotidiennement) "
        "consulter le babillard pour s'enquérir des dernières modifications."
    ).italic = True
    nb.paragraph_format.space_after = Pt(8)

    nb2 = doc.add_paragraph()
    nb2.add_run('2.').bold = True
    nb2.add_run(
        ' Les salles A, B, C, D, E de TP et Multimédia sont sur le site du '
        "Campus principal (Décanat de la Faculté des Sciences) et les salles "
        "F, G, H, I et J, K, L et O sur le site du campus secondaire (Lycée "
        "Classique). Quant aux salles M et N, elles se trouvent au campus C "
        "en face du CRA. Vous trouverez sur la porte de chaque salle la lettre "
        "correspondante. Et les salles telles que A.M, B.M, C.M, D.M, Amphi "
        "500, Terrain sont au campus de Monatélé."
    ).italic = True

    doc.add_paragraph()
    doyen = doc.add_paragraph()
    doyen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = doyen.add_run('Le Doyen')
    run.bold = True
    run.font.size = Pt(11)

    # ── Sortie en bytes ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    nom = (
        f'Planning_FS-UEB_'
        f'{semaine.date_debut:%Y-%m-%d}_au_{semaine.date_fin:%Y-%m-%d}.docx'
    )
    return buf.getvalue(), nom


# ── Sous-fonctions ──────────────────────────────────────────────────────────
def _ecrire_bloc_fr(cell, semaine):
    """Bloc texte français de l'en-tête."""
    lignes = [
        ('RÉPUBLIQUE DU CAMEROUN',              True),
        ('PAIX – TRAVAIL – PATRIE',             False),
        ('*******',                             False),
        ("MINISTÈRE DE L'ENSEIGNEMENT SUPÉRIEUR", True),
        ('*******',                             False),
        ("UNIVERSITÉ D'ÉBOLOWA",                True),
        ('*******',                             False),
        ('FACULTÉ DES SCIENCES',                True),
        ('*******',                             False),
        ('DIVISION DES AFFAIRES ACADÉMIQUES,',  True),
        ('DE LA SCOLARITÉ ET DE LA RECHERCHE',  True),
        ('*******',                             False),
    ]
    cell.text = ''  # vider
    for txt, bold in lignes:
        p = cell.add_paragraph()
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(8)

    ref = cell.add_paragraph()
    rr = ref.add_run(f'N° {semaine.numero_reference or "…………"}/UEb/DFS/DAARS')
    rr.bold = True
    rr.font.size = Pt(9)


def _ecrire_bloc_logos(cell):
    """Logos UEB + FS empilés dans la colonne centrale."""
    assets = Path(settings.BASE_DIR) / 'core' / 'static' / 'exports'
    cell.text = ''
    for nom_img in ('logo_ueb.png', 'logo_fs.png'):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chemin = assets / nom_img
        if chemin.exists():
            p.add_run().add_picture(str(chemin), width=Cm(2.5))


def _ecrire_bloc_en(cell):
    """Bloc texte anglais de l'en-tête."""
    lignes = [
        ('REPUBLIC OF CAMEROON',                                  True),
        ('PEACE-WORK-FATHERLAND',                                 False),
        ('******',                                                False),
        ('MINISTRY OF HIGHER EDUCATION',                          True),
        ('*******',                                               False),
        ('THE UNIVERSITY OF EBOLOWA',                             True),
        ('*******',                                               False),
        ('FACULTY OF SCIENCE',                                    True),
        ('*******',                                               False),
        ('DIVISION OF ACADEMIC AFFAIRS,',                         True),
        ('EDUCATION AND RESEARCH',                                True),
        ('*******',                                               False),
    ]
    cell.text = ''
    for txt, bold in lignes:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(8)

    ed = cell.add_paragraph()
    ed.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    er = ed.add_run(f'Ebolowa, le {date.today():%d/%m/%Y}')
    er.italic = True
    er.font.size = Pt(9)


def _construire_tableau(doc, seances_salle):
    """Tableau 7 colonnes × 8 lignes (1 header + 4 créneaux + 3 pauses)."""
    table = doc.add_table(rows=8, cols=7)
    table.style = 'Light Grid Accent 1'

    # En-têtes
    headers = ['Horaires'] + [j.label for j in Jour]
    for col, label in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True; r.font.size = Pt(9)

    # Lignes alternées créneau / pause (row_idx commence à 1)
    row_idx = 1
    for creneau_idx in range(4):
        deb, fin = HORAIRES_CRENEAUX[creneau_idx]
        row = table.rows[row_idx]

        # Colonne 0 : horaire cours
        cell = row.cells[0]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(_hum_horaire(deb, fin))
        r.bold = True; r.font.size = Pt(8)

        # Colonnes 1..6 : un par jour
        for jour_idx in range(6):
            cell = row.cells[jour_idx + 1]
            cell.text = ''
            s = next(
                (x for x in seances_salle if x.creneau == creneau_idx and x.jour == jour_idx),
                None,
            )
            if s is None:
                continue
            p = cell.paragraphs[0]
            r = p.add_run(f'{s.filiere.nom} {s.filiere.niveau}\n')
            r.bold = True; r.font.size = Pt(8)
            p.add_run(f'{s.ue.code}\n').font.size = Pt(8)
            p.add_run(f'{s.ue.intitule}\n').font.size = Pt(8)
            ens = f'{s.enseignant.get_grade_display()} {s.enseignant.nom}' if s.enseignant else '—'
            ens_run = p.add_run(ens)
            ens_run.italic = True
            ens_run.font.size = Pt(7.5)

        row_idx += 1

        # Ligne de pause après les 3 premiers créneaux
        if creneau_idx < 3:
            pd, pf = _PAUSES_HORAIRES[creneau_idx]
            pause_row = table.rows[row_idx]

            pcell = pause_row.cells[0]
            pcell.text = ''
            pp = pcell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr = pp.add_run(_hum_horaire(pd, pf))
            pr.italic = True; pr.font.size = Pt(7.5)
            _set_cell_bg(pcell, 'F3F4F6')

            for jour_idx in range(6):
                jcell = pause_row.cells[jour_idx + 1]
                jcell.text = ''
                _set_cell_bg(jcell, 'F3F4F6')

            row_idx += 1
